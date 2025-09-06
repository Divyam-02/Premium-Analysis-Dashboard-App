import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import streamlit as st
from io import BytesIO
from docx import Document
from docx.shared import Inches
import datetime

st.title("Premium Analysis Dashboard")

# ---- File uploader (CSV or Excel) ----
uploaded_file = st.file_uploader("Upload dataset (CSV or Excel)", type=["csv", "xlsx"])

# Initialize dictionary to store figures for Word report
if 'figures_dict' not in st.session_state:
    st.session_state.figures_dict = {}


def clean_premium_column(df, col):
    return (
        df[col]
        .astype(str)
        .str.replace(',', '', regex=False)
        .str.replace('EUR', '', regex=False)
        .astype(float)
    )

def analyze_premiums(df):
    # Convert date columns if available
    if "Policy Inception Date" in df.columns:
        df["Policy Inception Date"] = pd.to_datetime(df["Policy Inception Date"], errors="coerce")
    if "Policy Expiration Date" in df.columns:
        df["Policy Expiration Date"] = pd.to_datetime(df["Policy Expiration Date"], errors="coerce")

    # ---- Clean text columns ----
    df['Cedant'] = df['Cedant'].astype(str).str.strip()
    df['Responding LoB'] = df['Responding LoB'].astype(str).str.strip().str.lower()
    df['Broker*'] = df['Broker*'].astype(str).str.strip()

    # ---- Cedant Dropdown ----
    cedant_options = ["All"] + sorted(df['Cedant'].dropna().unique())
    selected_cedant = st.selectbox("Select Cedant", options=cedant_options)

    # ---- Filter dataset for selected Cedant ----
    if selected_cedant != "All":
        df_cedant = df[df['Cedant'] == selected_cedant]
    else:
        df_cedant = df.copy()

    # ---- LoB Dropdown (Cedant-dependent) ----
    lob_options = ["All"] + sorted(df_cedant['Responding LoB'].dropna().unique())
    selected_lob = st.selectbox("Select Line of Business (LoB)", lob_options)

    # ---- Broker Dropdown (Cedant + LoB dependent) ----
    if selected_lob != "All":
        df_for_brokers = df_cedant[df_cedant['Responding LoB'] == selected_lob]
    else:
        df_for_brokers = df_cedant.copy()
    broker_options = ["All"] + sorted(df_for_brokers['Broker*'].dropna().unique())
    selected_broker = st.selectbox("Select Broker", broker_options)

    # ---- Date filter selection (AFTER dropdowns) ----
    date_filter_type = st.selectbox(
        "Select Date Filter Type",
        [
            "Active Policies",
            "Policy Inception Date",
            "Policy Expiration Date",
            "No Date Filter"
        ],
        index=0  # default = Active Policies
    )

    inception_min = df["Policy Inception Date"].min() if "Policy Inception Date" in df.columns else None
    expiration_max = df["Policy Expiration Date"].max() if "Policy Expiration Date" in df.columns else None

    df_filtered = df_cedant.copy()  # base copy for filtering

    if date_filter_type == "Policy Inception Date":
        # Free manual selection (2000–2035 range)
        start_date = st.date_input(
            "Select Inception Start Date",
            value=pd.to_datetime("2000-01-01").date(),
            min_value=pd.to_datetime("2000-01-01").date(),
            max_value=pd.to_datetime("2035-12-31").date()
        )
        end_date = st.date_input(
            "Select Inception End Date",
            value=pd.to_datetime("2035-12-31").date(),
            min_value=pd.to_datetime("2000-01-01").date(),
            max_value=pd.to_datetime("2035-12-31").date()
        )

        df_filtered = df_filtered[
            (df_filtered["Policy Inception Date"] >= pd.to_datetime(start_date)) &
            (df_filtered["Policy Inception Date"] <= pd.to_datetime(end_date))
            ]

    elif date_filter_type == "Policy Expiration Date":
        # Free manual selection (2000–2035 range)
        start_date = st.date_input(
            "Select Expiration Start Date",
            value=pd.to_datetime("2000-01-01").date(),
            min_value=pd.to_datetime("2000-01-01").date(),
            max_value=pd.to_datetime("2035-12-31").date()
        )
        end_date = st.date_input(
            "Select Expiration End Date",
            value=pd.to_datetime("2035-12-31").date(),
            min_value=pd.to_datetime("2000-01-01").date(),
            max_value=pd.to_datetime("2035-12-31").date()
        )

        df_filtered = df_filtered[
            (df_filtered["Policy Expiration Date"] >= pd.to_datetime(start_date)) &
            (df_filtered["Policy Expiration Date"] <= pd.to_datetime(end_date))
            ]

    elif date_filter_type == "Active Policies":
        today = pd.to_datetime("today").normalize()

        # Get dataset bounds
        earliest_inception = df_filtered["Policy Inception Date"].min().date()
        latest_expiration = df_filtered["Policy Expiration Date"].max().date()

        # User-selectable ranges with defaults
        selected_start_date = st.date_input(
            "Select Start Date",
            value=earliest_inception,
            min_value=earliest_inception,
            max_value=latest_expiration
        )
        selected_end_date = st.date_input(
            "Select End Date",
            value=latest_expiration,
            min_value=earliest_inception,
            max_value=latest_expiration
        )

        # Reset button (placed just below inputs)
        if st.button("Reset Dates to Full Range"):
            selected_start_date = earliest_inception
            selected_end_date = latest_expiration

        # Apply active filter logic
        df_filtered = df_filtered[
            (df_filtered["Policy Inception Date"] <= today) &
            (df_filtered["Policy Expiration Date"] >= today) &
            (df_filtered["Policy Inception Date"] >= pd.to_datetime(selected_start_date)) &
            (df_filtered["Policy Expiration Date"] <= pd.to_datetime(selected_end_date))
            ]

    # ---- Apply LoB and Broker filters ----
    if selected_lob != "All":
        df_filtered = df_filtered[df_filtered['Responding LoB'] == selected_lob]
    if selected_broker != "All":
        df_filtered = df_filtered[df_filtered['Broker*'] == selected_broker]

    # ---- Display included/excluded counts ----
    st.subheader("Policy Count Overview")
    included_count = len(df_filtered)
    excluded_count = len(df_cedant) - included_count
    st.write(f"Included Policies: {included_count}")
    st.write(f"Excluded Policies: {excluded_count}")

    if df_filtered.empty:
        st.warning(f"No data found for Cedant: {selected_cedant}, LoB: {selected_lob}, Broker: {selected_broker}.")
        return

    # ---- Display All Filtered Policies ----
    st.subheader("Filtered Policy Details")

    # Full list of columns to display
    cols_to_show = [
        "Account Name", "Policy Inception Date", "Policy Expiration Date",
        "Submission Ref No.", "Responding LoB", "Responsible Underwriter",
        "Type", "Project Title - Policy Description", "Broker*", "Cedant",
        "AGCS Share %*", "AGCS Estimated Premium (full period) (converted)",
        "AGCS Estimated Premium (annual) (converted)",
        "AGCS GNWP (full period) - reporting (converted)",
        "AGCS GNWP (annual) - reporting (converted)",
        "Sub LoB", "Primary Risk Location", "Initiative Name",
        "AGCS Broker Segment", "Placement Source",
        "Additional Win / Loss Information", "Stage",
        "Class of Business - Corporate", "Reason Declined",
        "Multinational", "Lead"
    ]

    # Only keep the ones present in df_filtered
    available_cols = [c for c in cols_to_show if c in df_filtered.columns]

    if not df_filtered.empty:
        st.write(f"Total Policies: {len(df_filtered)}")
        st.dataframe(df_filtered[available_cols], use_container_width=True)
    else:
        st.warning("No policy data available for the selected filters.")

        # ---- Yearly Premiums Chart (NEW) ----
    if date_filter_type == "Policy Inception Date":
        df_filtered["Year"] = df_filtered["Policy Inception Date"].dt.year
    elif date_filter_type == "Policy Expiration Date":
        df_filtered["Year"] = df_filtered["Policy Expiration Date"].dt.year
    else:  # Active Policies OR No Date Filter
        df_filtered["Year"] = df_filtered["Policy Inception Date"].dt.year

        # Aggregate yearly premiums
    yearly_data = (
        df_filtered.groupby("Year")["AGCS GNWP (full period) - reporting (converted)"]
        .sum()
        .reset_index()
        .rename(columns={"AGCS GNWP (full period) - reporting (converted)": "Premium"})
    )
    yearly_data["Year"] = yearly_data["Year"].astype(int)  # ensure integer years

    st.subheader("Yearly Premium Trends")
    if not yearly_data.empty:
        fig, ax = plt.subplots(figsize=(10, 6))

        # Discrete x positions for bars
        x = np.arange(len(yearly_data))

        # --- NEW: Adjust bar width for single-year case ---
        bar_width = 0.4 if len(yearly_data) == 1 else 0.8
        if len(yearly_data) == 1:
            ax.set_xlim(-0.5, 0.5)  # optional: center the single bar

        # Bar plot
        bars = ax.bar(
            x,
            yearly_data["Premium"],
            width=bar_width,  # use dynamic width
            alpha=0.7,
            color="skyblue",
            label="Total Premium"
        )

        # Line plot connecting the bars
        ax.plot(
            x,
            yearly_data["Premium"],
            marker="o",
            color="#1f77b4",
            label="Growth Trend"
        )

        max_premium = yearly_data["Premium"].max()

        # Add exact premium labels on bars
        for bar in bars:
            height = bar.get_height()
            ax.text(
                bar.get_x() + bar.get_width() / 2,
                height,
                f"{height:,.0f}",
                ha="center",
                va="bottom",
                fontsize=9
            )

        # Add % growth labels slightly above line and shifted left
        for i in range(1, len(yearly_data)):
            prev = yearly_data.iloc[i - 1]["Premium"]
            curr = yearly_data.iloc[i]["Premium"]
            if prev > 0:
                pct_change = ((curr - prev) / prev) * 100
                ax.text(
                    x[i] - 0.15,
                    curr + 0.05 * max_premium,
                    f"{pct_change:+.1f}%",
                    ha="right",
                    va="bottom",
                    fontsize=8,
                    color="blue",
                    fontstyle="italic"
                )

        # Set x-axis labels
        ax.set_xticks(x)
        ax.set_xticklabels(yearly_data["Year"])

        ax.set_title(f"Yearly Premiums ({date_filter_type})", fontsize=14)
        ax.set_xlabel("Year")
        ax.set_ylabel("Total Premium (€) (in millions)")
        ax.legend()
        st.pyplot(fig)
        st.session_state.figures_dict["Yearly Premium Trends"] = fig

    else:
        st.warning("No yearly premium data available for the selected filters.")

    # --- Dedicated Yearly Premium Chart (Custom Year Ranges) ---
    st.subheader("Dedicated Yearly Premium Chart (For Comparison of Quarters year-wise)")

    # Get Cedant, Broker, LoB selections from the very first filters in the app
    df_custom = df.copy()

    # Apply only Cedant, Broker, LoB combination filter
    if selected_cedant != "All":
        df_custom = df_custom[df_custom["Cedant"] == selected_cedant]

    if selected_broker != "All":
        df_custom = df_custom[df_custom["Broker"] == selected_broker]

    if selected_lob != "All":
        df_custom = df_custom[df_custom["Responding LoB"] == selected_lob]

    # Proceed only if we still have data
    if not df_custom.empty:
        # Ask user how many years (custom ranges)
        st.markdown("### Select Number of Years to Compare")
        num_years = st.number_input(
            "Number of Years",
            min_value=1,
            max_value=10,
            value=3
        )

        yearly_premiums = []
        labels = []

        for i in range(num_years):
            st.markdown(f"**Year {i + 1} Selection**")

            # Start date input (wide range 2000–2035, default = dataset min)
            start_default = df_custom["Policy Inception Date"].min().date()
            start_date = st.date_input(
                f"Start Date for Year {i + 1}",
                min_value=pd.to_datetime("2000-01-01").date(),
                max_value=pd.to_datetime("2035-12-31").date(),
                value=min(max(start_default, pd.to_datetime("2000-01-01").date()), pd.to_datetime("2035-12-31").date()),
                key=f"start_{i}"
            )

            # Restrict end date to same year as start_date
            same_year_end = pd.Timestamp(start_date).replace(month=12, day=31).date()
            end_default = min(same_year_end, df_custom["Policy Inception Date"].max().date())
            end_date = st.date_input(
                f"End Date for Year {i + 1}",
                min_value=start_date,
                max_value=same_year_end,
                value=end_default if start_date <= end_default else start_date,  # keep within valid range
                key=f"end_{i}"
            )

            # Filter data for this custom range (Inception only)
            df_year = df_custom[
                (df_custom["Policy Inception Date"] >= pd.to_datetime(start_date)) &
                (df_custom["Policy Inception Date"] <= pd.to_datetime(end_date))
                ]

            total_premium = df_year["AGCS GNWP (full period) - reporting (converted)"].sum()
            yearly_premiums.append(total_premium)

            # Label with full custom range
            labels.append(f"{start_date.strftime('%d-%b-%Y')} → {end_date.strftime('%d-%b-%Y')}")

        # Plot the bar chart with connecting line
        if yearly_premiums:
            fig, ax = plt.subplots(figsize=(10, 6))

            bars = ax.bar(range(num_years), yearly_premiums, tick_label=labels, color="skyblue")

            # Add connecting line
            ax.plot(range(num_years), yearly_premiums, marker="o", color="blue", linestyle="-")

            # Add growth/decline labels
            max_premium = max(yearly_premiums) if yearly_premiums else 0
            for i in range(1, num_years):
                prev, curr = yearly_premiums[i - 1], yearly_premiums[i]
                if prev != 0:
                    pct_change = ((curr - prev) / prev) * 100
                    ax.text(
                        i - 0.15,
                        curr + 0.05 * max_premium,
                        f"{pct_change:+.1f}%",
                        ha="right",
                        va="bottom",
                        fontsize=8,
                        color="blue",
                        fontstyle="italic"
                    )

            ax.set_ylabel("Premium (Converted)")
            ax.set_title("Dedicated Yearly Premium Chart (For Comparison of Quarters year-wise)")

            # Rotate labels to avoid overlap
            plt.setp(ax.get_xticklabels(), rotation=30, ha="right")

            st.pyplot(fig)
            st.session_state.figures_dict["Dedicated Yearly Premium Chart (For Comparison of Quarters year-wise)"] = fig
    else:
        st.warning("No data available for the selected Cedant, Broker, and Responding LoB combination.")

    # ---- Premium Summaries ----
    total_premium_full = df_filtered['AGCS GNWP (full period) - reporting (converted)'].sum()
    total_premium_annual = df_filtered['AGCS GNWP (annual) - reporting (converted)'].sum()

    st.subheader("Premium Summary")
    col1, col2 = st.columns(2)
    col1.metric("Total Full Period Premium (€)", f"{total_premium_full:,.2f}")
    col2.metric("Total Annual Premium (€)", f"{total_premium_annual:,.2f}")

    # ---- Bar Chart ----
    fig, ax = plt.subplots(figsize=(6, 6))

    bar_width = 0.45  # narrower bars
    x = np.arange(2)  # positions of bars
    total_values = [total_premium_full, total_premium_annual]

    # Different shades of blue (2 bars)
    colors = plt.cm.Blues(np.linspace(0.5, 0.9, len(total_values)))

    bars = ax.bar(
        x,
        total_values,
        width=bar_width,
        color=colors
    )

    # Center the bars nicely on x-axis
    ax.set_xticks(x)
    ax.set_xticklabels(['Full Period', 'Annual'])
    ax.set_ylabel("Premium Amount (€)")
    ax.set_title(
        f"Total Premiums\nCedant: {selected_cedant}, LoB: {selected_lob}, Broker: {selected_broker}"
    )

    # Add exact values above bars
    for i, v in enumerate(total_values):
        ax.text(x[i], v, f"{v:,.0f}", ha='center', va='bottom', fontsize=8)

    # Optional: reduce space on sides of chart
    ax.margins(x=0.1)

    st.pyplot(fig)
    st.session_state.figures_dict["Total Premiums Bar Chart"] = fig

    # ---- LoB Pie Chart ----
    if selected_lob == "All":
        lob_summary = df_filtered.groupby('Responding LoB')[
            'AGCS GNWP (full period) - reporting (converted)'].sum().reset_index()
    else:
        total_premium_cedant = df_cedant['AGCS GNWP (full period) - reporting (converted)'].sum()
        lob_premium = df_filtered['AGCS GNWP (annual) - reporting (converted)'].sum()
        lob_summary = pd.DataFrame({
            'Responding LoB': [selected_lob, 'Other LoBs'],
            'AGCS GNWP (full period) - reporting (converted)': [
                lob_premium, max(total_premium_cedant - lob_premium, 0)
            ]
        })

    if lob_summary.empty or lob_summary['AGCS GNWP (full period) - reporting (converted)'].sum() <= 0:
        fig, ax = plt.subplots(figsize=(8, 8))
        ax.pie([1], labels=["No Premium Data"], autopct='%1.1f%%')
    else:
        fig, ax = plt.subplots(figsize=(8, 8))
        values = lob_summary['AGCS GNWP (full period) - reporting (converted)']
        labels = lob_summary['Responding LoB']

        #  Use shades of blue for all slices
        colors = plt.cm.Blues(np.linspace(0.4, 0.9, len(values)))

        wedges, texts, autotexts = ax.pie(
            values,
            labels=labels,  #  put LoB names back inside slices
            autopct=lambda p: f'{p:.1f}%\n({p * values.sum() / 100:,.0f}€)' if p > 0 else '',
            startangle=140,
            colors=colors
        )

        # Make text more readable
        for text in texts:
            text.set_fontsize(9)
        for autotext in autotexts:
            autotext.set_fontsize(8)
            autotext.set_color("white")  # so it contrasts with blue shades

        # 🔹 Build legend with premium values + %
        total_val = values.sum()
        legend_labels = [
            f"{l} – {v:,.0f} (€) – {v / total_val:.1%}"
            for l, v in zip(labels, values)
        ]

        ax.legend(
            wedges,
            legend_labels,
            title="LoBs",
            loc="upper center",
            bbox_to_anchor=(0.5, -0.1),  # place legend below pie
            fontsize=10,
            title_fontsize=12,
            ncol=2
        )

    ax.set_title(f"Premium Distribution by LoB\nCedant: {selected_cedant}")
    st.pyplot(fig)
    st.session_state.figures_dict["LoB Premium Distribution"] = fig

    # ---- Broker-level charts if selected broker = All ----
    if selected_broker == "All":
        broker_summary = df_filtered.groupby('Broker*').agg({
            'AGCS GNWP (full period) - reporting (converted)': 'sum',
            'AGCS GNWP (annual) - reporting (converted)': 'sum'
        }).reset_index()

        if not broker_summary.empty:
            x = np.arange(len(broker_summary))

            # ---- Broker-level charts if selected broker = All ----
            if selected_broker == "All":
                broker_summary = df_filtered.groupby('Broker*').agg({
                    'AGCS GNWP (full period) - reporting (converted)': 'sum',
                    'AGCS GNWP (annual) - reporting (converted)': 'sum'
                }).reset_index()

                if not broker_summary.empty:
                    x = np.arange(len(broker_summary))

                    # ---- Scatter + Line Plot ----
                    fig, ax = plt.subplots(figsize=(12, 6))
                    y_full = broker_summary['AGCS GNWP (full period) - reporting (converted)']
                    y_annual = broker_summary['AGCS GNWP (annual) - reporting (converted)']

                    # 🔹 Two shades of blue for consistency
                    full_color = '#003f5c' # deep navy blue
                    annual_color = '#6baed6'  # lighter blue

                    ax.scatter(x, y_full, color=full_color, s=60, label='Full Period')
                    ax.scatter(x, y_annual, color=annual_color, s=60, label='Annual')
                    ax.plot(x, y_full, '--', color=full_color, alpha=0.7)
                    ax.plot(x, y_annual, '--', color=annual_color, alpha=0.7)

                    # Annotate points
                    for i, v in enumerate(y_full):
                        ax.annotate(f"{v:,.0f}", (x[i], v), textcoords="offset points", xytext=(0, 5),
                                    ha='center', fontsize=8, color=full_color)
                    for i, v in enumerate(y_annual):
                        ax.annotate(f"{v:,.0f}", (x[i], v), textcoords="offset points", xytext=(0, -12),
                                    ha='center', fontsize=8, color=annual_color)

                    ax.set_xticks(x)
                    ax.set_xticklabels(broker_summary['Broker*'], rotation=45, ha='right')
                    ax.set_ylabel("Premium Amount (€)")
                    ax.set_title(f"Scatter + Line Plot of Brokers\nCedant: {selected_cedant}, LoB: {selected_lob}")
                    ax.legend()
                    st.pyplot(fig)
                    st.session_state.figures_dict["Broker Scatter + Line Plot"] = fig

            # ---- Bar Chart per Broker ----
            bar_width = 0.35
            fig, ax = plt.subplots(figsize=(12, 6))

            # 🔹 Two shades of blue
            full_color = '#1f77b4'  # standard blue
            annual_color = '#6baed6'  # lighter blue

            bars1 = ax.bar(x - bar_width / 2, y_full, width=bar_width, color=full_color, label='Full Period')
            bars2 = ax.bar(x + bar_width / 2, y_annual, width=bar_width, color=annual_color, label='Annual')

            # Add exact values
            for bar in bars1:
                ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height(), f"{bar.get_height():,.0f}",
                        ha='center', va='bottom', fontsize=8, color=full_color)
            for bar in bars2:
                ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height(), f"{bar.get_height():,.0f}",
                        ha='center', va='bottom', fontsize=8, color=annual_color)

            ax.set_xticks(x)
            ax.set_xticklabels(broker_summary['Broker*'], rotation=45, ha='right')
            ax.set_ylabel("Premium Amount (€)")
            ax.set_title(f"Bar Chart of Premiums per Broker\nCedant: {selected_cedant}, LoB: {selected_lob}")
            ax.legend()
            st.pyplot(fig)
            st.session_state.figures_dict["Broker Premium Bar Chart"] = fig

            # ---- Top 5 Brokers Pie Chart ----
            premium_type = "AGCS GNWP (full period) - reporting (converted)"
            cedant_broker_summary = df_filtered.groupby('Broker*')[premium_type].sum().reset_index()

            if not cedant_broker_summary.empty:
                top_brokers = cedant_broker_summary.sort_values(
                    by=premium_type, ascending=False
                ).head(5)

                # Ensure always 5 rows even if fewer brokers
                while len(top_brokers) < 5:
                    top_brokers = pd.concat([top_brokers, pd.DataFrame({
                        "Broker*": [f"No Data {len(top_brokers) + 1}"],
                        premium_type: [0]
                    })], ignore_index=True)

                # Add "Others" slice
                others_premium = max(
                    cedant_broker_summary[premium_type].sum() - top_brokers[premium_type].sum(),
                    0
                )

                pie_labels = list(top_brokers['Broker*']) + ['Others']
                pie_values = list(top_brokers[premium_type]) + [others_premium]

                fig, ax = plt.subplots(figsize=(8, 8))

                # 🔹 Dynamic shades of blue
                colors = plt.cm.Blues(np.linspace(0.4, 0.9, len(pie_values)))

                wedges, texts, autotexts = ax.pie(
                    pie_values,
                    labels=pie_labels,
                    autopct=lambda p: f'{p:.1f}%\n({sum(pie_values) * p / 100:,.0f})',
                    startangle=140,
                    colors=colors
                )

                # Make text more readable
                for text in texts:
                    text.set_fontsize(9)
                for autotext in autotexts:
                    autotext.set_fontsize(8)
                    autotext.set_color("white")  # so it contrasts with blue shades

                # 🔹 Dynamic legend
                total_val = sum(pie_values)
                legend_labels = [f"{l} – {v:,.0f} (€) – {v / total_val:.1%}" for l, v in zip(pie_labels, pie_values)]

                ax.legend(
                    wedges,
                    legend_labels,
                    title="Brokers",
                    loc="upper center",
                    bbox_to_anchor=(0.5, -0.1),  # centers legend below pie
                    fontsize=10,
                    title_fontsize=12,
                    ncol=2
                )

                ax.set_title(f"Top 5 Brokers Premium Share\nCedant: {selected_cedant}\nPremium Type: {premium_type}")
                st.pyplot(fig)
                st.session_state.figures_dict["Top 5 Brokers Premium Share"] = fig

            else:
                fig, ax = plt.subplots(figsize=(8, 8))
                ax.pie([1], labels=["No Broker Data"], autopct='%1.1f%%', colors=["#87CEEB"])  # light blue fallback
                ax.set_title(f"Top 5 Brokers Premium Share\nCedant: {selected_cedant}")
                st.pyplot(fig)
                st.session_state.figures_dict["Top 5 Brokers Premium Share"] = fig

    def generate_word_report():
        doc = Document()
        doc.add_heading("Premium Analysis Report", level=0)
        doc.add_paragraph(f"Generated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph("")

        # Filters
        doc.add_heading("Filters Applied", level=1)
        doc.add_paragraph(f"Cedant: {selected_cedant}")
        doc.add_paragraph(f"LoB: {selected_lob}")
        doc.add_paragraph(f"Broker: {selected_broker}")
        doc.add_paragraph(f"Date Filter Type: {date_filter_type}")
        if date_filter_type in ["Policy Inception Date", "Policy Expiration Date"]:
            doc.add_paragraph(f"Date Range: {start_date} → {end_date}")
        elif date_filter_type == "Active Policies":
            doc.add_paragraph(f"Active Policy Range: {selected_start_date} → {selected_end_date}")
        else:
            doc.add_paragraph("Date Range: None")

        # Premium Summary
        doc.add_heading("Premium Summary", level=1)
        doc.add_paragraph(f"Total Full Period Premium (€): {total_premium_full:,.2f}")
        doc.add_paragraph(f"Total Annual Premium (€): {total_premium_annual:,.2f}")

        # Add all figures
        for title, fig in st.session_state.figures_dict.items():
            img_stream = BytesIO()
            fig.savefig(img_stream, format='png', bbox_inches='tight')
            img_stream.seek(0)
            doc.add_heading(title, level=1)
            doc.add_picture(img_stream, width=Inches(6))

        # Save to memory
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    word_report = generate_word_report()

    st.download_button(
        label=" Download Word Report ",
        data=word_report,
        file_name="premium_analysis_report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


# ---- File upload handler ----
if uploaded_file:
    if uploaded_file.name.lower().endswith(".csv"):
        df = pd.read_csv(uploaded_file)
    elif uploaded_file.name.lower().endswith(".xlsx"):
        df = pd.read_excel(uploaded_file, engine="openpyxl")
    else:
        st.error("Please upload only CSV or XLSX files.")
        st.stop()

    #  Clean column names to avoid dash issues
    df.columns = df.columns.str.strip()
    df.columns = df.columns.str.replace("–", "-", regex=False)  # en-dash → hyphen
    df.columns = df.columns.str.replace("\u2013", "-", regex=False)  # unicode en-dash
    df.columns = df.columns.str.replace("\u2014", "-", regex=False)  # em-dash

    for col in ['AGCS GNWP (full period) - reporting (converted)', 'AGCS GNWP (annual) - reporting (converted)']:
        if col in df.columns:
            df[col] = clean_premium_column(df, col)

    analyze_premiums(df)
else:
    st.info("Please upload a CSV or Excel file to continue.")
